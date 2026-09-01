from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "PAPER_DRAFT_CN.md"
OUTPUT = ROOT / "PAPER_DRAFT_CN.docx"


def set_font(run, name="宋体", size=11, bold=False, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("RockSeg | PAPER_DRAFT_CN | ")
    set_font(run, "Arial", 8, color=(110, 110, 110))
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, (31, 78, 121), 16, 8),
        ("Heading 2", 13, (31, 78, 121), 12, 6),
        ("Heading 3", 11.5, (55, 55, 55), 10, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Formula" not in [s.name for s in doc.styles]:
        formula = doc.styles.add_style("Formula", WD_STYLE_TYPE.PARAGRAPH)
    else:
        formula = doc.styles["Formula"]
    formula.font.name = "Cambria Math"
    formula._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    formula.font.size = Pt(10.5)
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(5)
    formula.paragraph_format.space_after = Pt(7)

    add_page_number(section.footer.paragraphs[0])


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("RockSeg 论文中文初稿")
    set_font(r, "黑体", 20, bold=True, color=(31, 78, 121))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("当前版本：方法章节（Chapter 3）")
    set_font(r, "宋体", 11, color=(90, 90, 90))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Status: DRAFT -- Chinese-first manuscript based on frozen experimental evidence")
    set_font(r, "Arial", 9, color=(110, 110, 110))

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(14)
    note.paragraph_format.left_indent = Cm(0.35)
    note.paragraph_format.right_indent = Cm(0.35)
    r = note.add_run("文档范围说明：当前 Word 版本依据 PAPER_DRAFT_CN.md 生成，仅包含第 3 章方法正文（3.1--3.6）。第 1、2、4、5、6 章及摘要、参考文献尚未写入当前正文稿。")
    set_font(r, "宋体", 10, color=(90, 90, 90))


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.add_run(text)
    for run in p.runs:
        set_font(run, "宋体", 11)
    return p


def add_list_item(doc, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    for run in p.runs:
        set_font(run, "宋体", 11)
    return p


def add_formula(doc, formula_lines):
    text = " ".join(line.strip() for line in formula_lines).strip()
    p = doc.add_paragraph(style="Formula")
    r = p.add_run(text)
    set_font(r, "Cambria Math", 10.5)
    return p


def build():
    source = SOURCE.read_text(encoding="utf-8")
    lines = source.splitlines()
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    in_formula = False
    formula_lines = []
    status_skipped = False
    for raw in lines:
        line = raw.rstrip()
        if not status_skipped and line.startswith("状态："):
            status_skipped = True
            continue
        if line.strip() == "$$":
            if in_formula:
                add_formula(doc, formula_lines)
                formula_lines = []
                in_formula = False
            else:
                in_formula = True
            continue
        if in_formula:
            formula_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:].strip(), style="Heading 1")
            for run in p.runs:
                set_font(run, "黑体", 16, bold=True, color=(31, 78, 121))
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip(), style="Heading 2")
            for run in p.runs:
                set_font(run, "黑体", 13, bold=True, color=(31, 78, 121))
        elif line.startswith("#### "):
            p = doc.add_paragraph(line[5:].strip(), style="Heading 3")
            for run in p.runs:
                set_font(run, "黑体", 11.5, bold=True, color=(55, 55, 55))
        elif re.match(r"^\d+\.\s+", line):
            add_list_item(doc, re.sub(r"^\d+\.\s+", "", line), numbered=True)
        elif line.startswith("- "):
            add_list_item(doc, line[2:].strip())
        else:
            add_paragraph(doc, line)

    doc.core_properties.title = "RockSeg 论文中文初稿 - 方法章节"
    doc.core_properties.subject = "Frozen evidence-based manuscript draft"
    doc.core_properties.author = "RockSeg research team"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
