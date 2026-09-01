from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = Path(
    r"C:\Users\Administrator\WPSDrive\1714584739\WPS企业云盘\杭州电子科技大学\我的企业文档\Measurement期刊\中文论文第二章_材料与方法与第三章结果框架.docx"
)
A_RESULTS = PROJECT_ROOT / "docs" / "results" / "current_results.json"
B_RESULTS = PROJECT_ROOT / "experiments" / "site_b_run" / "outputs" / "site_b_quality_check_summary.json"


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def fmt(value, digits: int = 4) -> str:
    if isinstance(value, int):
        return str(value)
    if isinstance(value, float):
        text = f"{value:.{digits}f}"
        return text.rstrip("0").rstrip(".")
    return str(value)


def set_run_font(paragraph, east_asia: str = "宋体", ascii_font: str = "Times New Roman") -> None:
    for run in paragraph.runs:
        run.font.name = ascii_font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
        run.font.size = Pt(10.5)


def qn(tag: str):
    from docx.oxml.ns import qn as _qn

    return _qn(tag)


def delete_from_heading3(document: Document) -> None:
    body = document._body._element
    children = list(body)
    start_idx = None
    for idx, child in enumerate(children):
        if not child.tag.endswith("}p"):
            continue
        para = None
        for p in document.paragraphs:
            if p._p is child:
                para = p
                break
        if para is None:
            continue
        text = para.text.strip()
        if para.style.name == "Heading 1" and text.startswith("3"):
            start_idx = idx
            break
    if start_idx is None:
        raise RuntimeError("Could not find the chapter 3 heading.")

    for child in children[start_idx:]:
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def add_paragraph(document: Document, text: str, style: str = "Normal"):
    p = document.add_paragraph(text, style=style)
    set_run_font(p)
    return p


def add_caption(document: Document, text: str):
    p = document.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p)
    return p


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(paragraph)


def main() -> None:
    a = load_json(A_RESULTS)
    b = load_json(B_RESULTS)
    bk = b["key_results"]

    doc = Document(DOCX_PATH)
    delete_from_heading3(doc)

    add_paragraph(doc, "3 结果", "Heading 1")
    add_paragraph(
        doc,
        "本章基于前述固定流程，对矿区 A 的基线结果与矿区 B 的独立运行结果进行整理和分析。"
        "矿区 B 在不重新训练分割模型、不改变主要处理流程的条件下完成运行，因此其结果主要用于考察该 DOM-点云协同测量流程在另一处矿区场景中的可执行性和输出稳定性。"
        "需要说明的是，本章中的三维几何筛查保留率表示候选目标经过点云几何约束后的保留比例，并不等同于人工标注意义上的检测精度；二维代理体积与 2.5D 体积之间的差异用于说明估算方法的差别，也不作为真实体积误差。",
    )

    add_paragraph(doc, "3.1 跨矿区流程运行结果", "Heading 2")
    add_paragraph(
        doc,
        "两个矿区均完成了从自适应切片、实例分割、重叠切片融合、三维几何筛查到 2.5D 体积估算的完整流程。"
        f"矿区 A 的 DOM 覆盖面积为 {fmt(a['scene']['dom']['area_m2'])} m²，四叉树切片后保留 {a['slicing']['quadtree_dom']['kept_tiles']} 个有效切片；"
        f"矿区 B 的 DOM 覆盖面积为 {fmt(bk['dom_area_m2'])} m²，保留 {bk['slicing_kept_tiles']} 个有效切片。"
        "两处场景的有效切片数量并不完全随 DOM 面积线性变化，这与石料堆积密度、影像边缘分布以及无效背景比例有关。",
    )
    add_paragraph(
        doc,
        f"在固定分割模型和直径过滤约束下，矿区 A 得到 {a['detection']['diameter_filtered_detections']} 个二维候选检测结果，"
        f"矿区 B 得到 {bk['detection_count']} 个二维候选检测结果。"
        "随后，同一融合策略将重叠切片中的重复检测合并为逐石候选对象。"
        f"矿区 A 的融合候选数为 {a['fusion']['candidate_stones']}，矿区 B 为 {bk['p1_candidate_stones']}。"
        "该结果表明，切片重叠带来的重复目标在两个矿区中均可以被合并到逐石尺度，为后续三维筛查和体积估算提供统一对象单元。",
    )
    add_caption(doc, "表 3-1 两个矿区的流程输出规模")
    add_table(
        doc,
        ["指标", "矿区 A", "矿区 B"],
        [
            ["DOM 面积/m²", fmt(a["scene"]["dom"]["area_m2"]), fmt(bk["dom_area_m2"])],
            ["有效切片数", str(a["slicing"]["quadtree_dom"]["kept_tiles"]), str(bk["slicing_kept_tiles"])],
            ["二维检测数", str(a["detection"]["diameter_filtered_detections"]), str(bk["detection_count"])],
            ["融合候选石块数", str(a["fusion"]["candidate_stones"]), str(bk["p1_candidate_stones"])],
            ["三维筛查后保留数", str(a["fusion"]["accepted_stones"]), str(bk["p1_accepted_stones"])],
            ["三维筛查剔除数", str(a["fusion"]["rejected_stones"]), str(bk["p1_rejected_stones"])],
            ["2.5D 体积估算对象数", str(a["volume"]["qc_passed"]), str(bk["volume_qc_passed"])],
        ],
    )

    add_paragraph(doc, "3.2 三维几何筛查对候选对象的影响", "Heading 2")
    add_paragraph(
        doc,
        "为了区分仅由二维影像分割得到的候选对象与具有点云几何支撑的石块对象，本研究设置了不启用三维筛查的 P0 流程和启用三维筛查的 P1 流程。"
        f"在矿区 B 中，P0 将 {bk['detection_count']} 个二维检测结果融合为 {bk['p0_candidate_stones']} 个候选对象，且不进行几何剔除；"
        f"P1 使用相同的二维检测与融合候选，经过三维几何筛查后保留 {bk['p1_accepted_stones']} 个对象，剔除 {bk['p1_rejected_stones']} 个对象。"
        f"对应的筛查保留率为 {fmt(bk['p1_acceptance_ratio'] * 100, 2)}%，剔除率为 {fmt(bk['p1_rejection_ratio'] * 100, 2)}%。",
    )
    reasons = bk["p1_rejection_reasons"]
    add_paragraph(
        doc,
        "从剔除原因看，矿区 B 中主要问题集中在相对高度分位数不足和抬升点比例不足，分别记录为"
        f" {reasons.get('insufficient_p90_height', 0)} 次和 {reasons.get('insufficient_elevated_ratio', 0)} 次；"
        f"局部高程范围不足记录为 {reasons.get('insufficient_z_range', 0)} 次，点数不足记录为 {reasons.get('too_few_points', 0)} 次。"
        "由于同一对象可能同时触发多个筛查条件，各类原因的次数之和不要求等于剔除对象总数。"
        "这一结果说明三维筛查主要发挥的是几何一致性约束作用：它可以排除一部分在二维影像中形态相似、但缺少足够三维起伏或点云支撑的候选对象。",
    )
    add_caption(doc, "表 3-2 矿区 B 中 P0 与 P1 流程结果对比")
    add_table(
        doc,
        ["流程", "输入二维检测数", "融合候选数", "三维筛查", "保留对象数", "剔除对象数"],
        [
            ["P0", str(bk["detection_count"]), str(bk["p0_candidate_stones"]), "未启用", str(bk["p0_accepted_stones"]), "0"],
            ["P1", str(bk["detection_count"]), str(bk["p1_candidate_stones"]), "启用", str(bk["p1_accepted_stones"]), str(bk["p1_rejected_stones"])],
        ],
    )

    add_paragraph(doc, "3.3 逐石 2.5D 体积估算结果", "Heading 2")
    add_paragraph(
        doc,
        f"通过三维筛查后，矿区 A 有 {a['volume']['qc_passed']} 个对象进入 2.5D 体积统计，"
        f"矿区 B 有 {bk['volume_qc_passed']} 个对象进入统计。"
        f"矿区 A 的 2.5D 总体积为 {fmt(a['volume']['volume_2d5_m3']['sum'])} m³，"
        f"均值为 {fmt(a['volume']['volume_2d5_m3']['mean'])} m³，中位数为 {fmt(a['volume']['volume_2d5_m3']['median'])} m³；"
        f"矿区 B 的 2.5D 总体积为 {fmt(bk['volume_2d5_sum_m3'])} m³，"
        f"均值为 {fmt(bk['volume_2d5_mean_m3'])} m³，中位数为 {fmt(bk['volume_2d5_median_m3'])} m³。"
        "两处矿区的单石体积均呈现均值高于中位数的特征，说明体积分布受到少量较大石块的影响。",
    )
    add_caption(doc, "表 3-3 两个矿区的体积估算结果")
    add_table(
        doc,
        ["指标", "矿区 A", "矿区 B"],
        [
            ["体积统计对象数", str(a["volume"]["qc_passed"]), str(bk["volume_qc_passed"])],
            ["2.5D 总体积/m³", fmt(a["volume"]["volume_2d5_m3"]["sum"]), fmt(bk["volume_2d5_sum_m3"])],
            ["2.5D 平均体积/m³", fmt(a["volume"]["volume_2d5_m3"]["mean"]), fmt(bk["volume_2d5_mean_m3"])],
            ["2.5D 中位体积/m³", fmt(a["volume"]["volume_2d5_m3"]["median"]), fmt(bk["volume_2d5_median_m3"])],
            ["二维代理总体积/m³", fmt(a["volume"]["volume_2d_proxy_m3"]["sum"]), fmt(bk["volume_2d_proxy_sum_m3"])],
            ["二维代理与 2.5D 相关系数", fmt(a["volume"]["comparison"]["pearson_r"]), fmt(bk["proxy_2d5_pearson_r"])],
        ],
    )

    add_paragraph(doc, "3.4 不同粒径区间的数量与体积贡献", "Heading 2")
    add_paragraph(
        doc,
        "进一步按照等效直径将矿区 B 的保留对象划分为四个粒径区间。"
        "结果显示，0.50-0.75 m 区间的石块数量占比最高，说明小粒径石块在数量上占据主体；"
        "但随着粒径增大，单石平均体积明显增加，较大粒径区间虽然数量占比较低，却仍对总体积贡献较高。"
        "这种数量贡献与体积贡献之间的不一致，是逐石体积估算相较于单纯数量统计更有意义的原因之一。",
    )
    add_caption(doc, "表 3-4 矿区 B 不同粒径区间的数量与体积贡献")
    add_table(
        doc,
        ["等效直径区间/m", "石块数", "数量占比/%", "2.5D 体积/m³", "体积占比/%", "平均体积/m³"],
        [
            [
                row["bin_label"],
                str(row["stone_count"]),
                fmt(row["count_ratio"] * 100, 2),
                fmt(row["volume_2d5_sum_m3"]),
                fmt(row["volume_2d5_ratio"] * 100, 2),
                fmt(row["volume_2d5_mean_m3"]),
            ]
            for row in bk["diameter_bins"]
        ],
    )

    add_paragraph(doc, "3.5 二维代理体积与 2.5D 体积的对比", "Heading 2")
    add_paragraph(
        doc,
        f"矿区 A 的二维代理总体积为 {fmt(a['volume']['volume_2d_proxy_m3']['sum'])} m³，"
        f"高于对应的 2.5D 总体积 {fmt(a['volume']['volume_2d5_m3']['sum'])} m³；"
        f"矿区 B 的二维代理总体积为 {fmt(bk['volume_2d_proxy_sum_m3'])} m³，"
        f"同样高于 2.5D 总体积 {fmt(bk['volume_2d5_sum_m3'])} m³。"
        f"在逐石尺度上，矿区 A 中两种体积指标的 Pearson 相关系数为 {fmt(a['volume']['comparison']['pearson_r'])}，"
        f"矿区 B 为 {fmt(bk['proxy_2d5_pearson_r'])}。"
        "这说明二维尺度信息与 2.5D 体积之间存在较强相关性，但二维代理方法倾向于将等效直径直接映射为规则几何体，无法反映石块可见表面的真实起伏和局部高度变化。",
    )
    add_paragraph(
        doc,
        "因此，二维代理体积更适合作为方法对照，而不是最终测量结果。"
        "在矿山现场的石料堆积场景中，石块外形不规则、遮挡和接触关系复杂，仅依赖二维轮廓容易放大或压缩部分对象的体积估计。"
        "引入点云高度信息后，2.5D 方法能够在同一二维边界内利用局部 GroundDEM 作为参考面，对可见表面的高度积分进行估算，从而更符合本文面向工程测量流程的研究目标。",
    )

    add_paragraph(doc, "3.6 小结", "Heading 2")
    add_paragraph(
        doc,
        "总体来看，固定流程在两个矿区均完成了从影像分割到逐石体积估算的自动化运行，矿区 B 的独立运行结果与矿区 A 基线结果在流程闭合性上保持一致。"
        "三维几何筛查在不改变二维分割输入的情况下，对缺少点云几何支撑的候选对象进行了进一步约束；"
        "2.5D 体积统计则提供了比二维数量和二维代理体积更接近三维测量需求的结果形式。"
        "不过，当前结果仍应理解为流程输出和内部对照结果。若要进一步报告检测精度或绝对体积精度，还需要补充人工标注样本、现场量测或其他可靠参考数据。",
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
