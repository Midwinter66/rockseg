from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from build_chinese_paper_framework import (
    add_formula, add_heading, add_para, add_placeholder, add_rich_para,
    math_fraction, math_radical, math_run, math_sub, math_sup,
    math_sup_group, remove_existing_body,
)


PAPER_ROOT = Path(r"C:\Users\Administrator\WPSDrive\1714584739\WPS企业云盘\杭州电子科技大学\我的企业文档\Measurement期刊")
OUTPUT = PAPER_ROOT / "中文论文第二章_材料与方法与第三章结果框架.docx"
LOCKED_OUTPUT_FALLBACK = PAPER_ROOT / "中文论文第二章_材料与方法与第三章结果框架_Fig2修订版.docx"
SOURCE_CANDIDATES = [
    PAPER_ROOT / "中文论文框架_修订版.docx",
    OUTPUT,
]


def resolve_source_docx() -> Path:
    for path in SOURCE_CANDIDATES:
        if path.exists():
            return path
    candidates = "；".join(str(path) for path in SOURCE_CANDIDATES)
    raise FileNotFoundError(f"未找到可用的 DOCX 源文件：{candidates}")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, centered: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_table(document, title: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(3)
    run = caption.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Cm(widths[i])
        set_cell_text(cell, value, bold=True, centered=True)
        set_cell_shading(cell, "D9EAF7")
    prevent_row_split(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].width = Cm(widths[i])
            set_cell_text(cells[i], value, centered=(i == 0))
        prevent_row_split(table.rows[-1])
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.15)
    section.right_margin = Cm(2.15)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def build_document() -> Path:
    document = Document(resolve_source_docx())
    remove_existing_body(document)
    configure_document(document)

    add_heading(document, "2 材料与方法", 1)
    add_heading(document, "2.1 研究区与输入数据", 2)
    add_para(document, "本研究以两个露天矿岩石堆积场景为对象，构建并验证 DOM 与点云协同的岩块量测流程。每个场景均由高分辨率数字正射影像（digital orthophoto map, DOM）及同一倾斜摄影重建成果导出的表面点云组成。DOM 提供岩块边界、纹理和投影尺度，点云提供可见上表面的三维坐标。需要说明的是，本文所用点云由 OSGB 摄影测量模型转换获得，并非 LiDAR 点云。")
    add_para(document, "矿区 A 用于流程开发、参数确定和冻结；矿区 B 在不改变模型权重与处理参数的条件下用于独立运行。图 2-1 给出了两个研究区的输入数据对照，其中图 2-1(a,b) 为矿区 A 的 DOM 与同一坐标窗口内摄影测量点云，图 2-1(c,d) 为矿区 B 的 DOM 与同一坐标窗口内摄影测量点云。")
    add_para(document, "两类数据进入计算前均进行空间一致性检查：DOM 的世界文件用于建立像素与投影坐标的关系，点云需转换至同一平面参考框架。该检查是后续掩膜回投、跨切片融合和局部点云裁取能够对应同一岩块的前提。")
    add_formula(document, [math_run("x = C + Au + Bv,     y = F + Du + Ev")], 1)
    add_para(document, "式中，(u, v) 为像素坐标，(x, y) 为对应的投影坐标；A 和 E 表示像元尺度，B 和 D 表示旋转项，C 和 F 表示影像原点的投影坐标。两个场景的 DOM 空间分辨率均为 0.01 m。矿区 B 在正式独立运行前需完成点云局部坐标与 DOM 投影坐标之间转换关系的核验，并将该转换与实验配置一并保存。")
    add_table(document, "表 2-1 研究区输入数据概况", ["研究区", "DOM", "摄影测量表面点云", "在本文中的作用"], [
        ["矿区 A", "8783 × 21713 像素\n空间分辨率：0.01 m\n覆盖面积：19070.53 m²\n坐标参考：EPSG:4536", "两个 LAZ 数据块\n总点数：146721392\n绝对世界坐标", "用于方法开发、参数确定和流程冻结"],
        ["矿区 B", "7269 × 22729 像素\n空间分辨率：0.01 m\n覆盖面积：16521.71 m²\n坐标参考：EPSG:4536", "两个 LAZ 数据块\n总点数：187360460\n局部坐标，需完成转换核验", "使用冻结流程进行独立运行"],
    ], [2.0, 5.0, 4.7, 4.9])
    add_placeholder(document, "此处插入图 2-1", "图 2-1 两个研究区 DOM 与摄影测量点云输入对照")

    add_heading(document, "2.2 预训练实例分割模型与数据独立性", 2)
    add_para(document, "岩块候选由预训练的 YOLO11m-seg 实例分割模型生成。模型对每个有效切片输出岩块掩膜、置信度和包围框。模型训练样本来自其他矿区，矿区 A 和矿区 B 均未参与训练或微调。因此，两个研究区中的模型推理不依赖场景内再训练，本文关注的是固定二维候选生成器与三维量测链条在不同矿区中的执行表现。")
    add_para(document, "为保持实验边界清晰，模型权重、输入尺寸和置信度阈值在矿区 A 完成参数确定后固定，矿区 B 不再据其运行结果调整。投稿前需补充模型训练数据的标注规范、训练集与验证集划分、数据增强、训练轮数、优化器和学习率等复现信息；这些信息用于说明模型来源，不应与本文两个研究区的独立运行数据混写。")

    add_heading(document, "2.3 DOM 与点云协同岩块测量流程", 2)
    add_para(document, "在输入数据完成空间一致性检查后，测量流程由 DOM 自适应切片、实例分割与二维量测、跨切片重复检测融合、点云几何筛查及 GroundDEM 参考的 2.5D 体积估计组成。各阶段以世界坐标为共同索引，最终为每个岩块输出可追溯的二维属性、三维筛查状态和体积估计。图 2-2 概括了从输入数据到逐石量测结果的主线流程。")
    add_placeholder(document, "此处插入图 2-2", "图 2-2 DOM 与点云协同岩块测量流程")

    add_heading(document, "2.3.1 自适应切片、实例分割与二维量测", 3)
    add_para(document, "大范围高分辨率 DOM 不能在保持原始地面分辨率的条件下直接输入实例分割网络。若将整幅影像缩放，岩块边界和小尺度纹理会被压缩；若采用统一小窗口遍历，则低纹理区域和影像外部背景同样产生推理开销。为此，本文以局部边缘分布为依据递归划分 DOM，使切片尺度随局部结构复杂度变化。")
    add_para(document, "DOM 首先转换为灰度影像并提取二值边缘图。对于候选区域 R，边缘密度定义为区域内边缘像素数占总像素数的比例：")
    add_formula(document, [math_sub("ρ", "e"), math_run("(R) = "), math_fraction([math_sub("N", "e"), math_run("(R)")], [math_sub("N", "p"), math_run("(R)")])], 2)
    add_para(document, "当区域内容不足或不含边缘时，该区域不进入推理；当边缘密度达到分裂条件且区域尺度仍大于最小切片尺度时，区域继续划分为四个子区域。相邻切片保留重叠带，以降低岩块跨越边界时的截断影响。每个保留切片记录其全局像素范围、世界坐标边界和处理状态，使局部检测结果能够恢复至全幅 DOM。")
    add_para(document, "保留切片输入 YOLO11m-seg 模型后，实例掩膜被用于计算二维投影面积和等效粒径。二维面积由掩膜中的前景像素总数及两个方向的地面采样距离共同确定；等效粒径则由等面积圆计算：")
    add_formula(document, [math_sub("A", "2D"), math_run(" = "), math_sub("n", "p"), math_sub("s", "x"), math_sub("s", "y"), math_run(",     "), math_sub("d", "eq"), math_run(" = "), math_radical([math_fraction([math_run("4"), math_sub("A", "2D")], [math_run("π")])])], 3)
    add_para(document, "等效粒径表示与掩膜投影面积相同的圆直径，用于对不规则岩块轮廓给出统一的一维尺度描述。本文将最小报告粒径作为统计范围约束；低于该尺度的候选不参与后续主分析。该阈值是报告协议的一部分，不等同于模型的理论最小检测能力。")

    add_heading(document, "2.3.2 跨切片重复检测融合", 3)
    add_para(document, "切片重叠保留了边界附近岩块的完整信息，也使同一岩块可能在不同切片中产生重复检测。为避免重复计数，掩膜质心和包围框先由式（1）转换至世界坐标，再在来自不同切片的检测之间建立关联。关联同时考虑质心距离、包围框交并比及来源切片约束；满足条件的候选通过相关聚类合并为单一岩块对象。该过程对应图 2-2 中的跨切片重复检测融合环节。")
    add_formula(document, [math_sub("w", "ij"), math_run(" = exp[−"), math_sup_group([math_run("||"), math_sub("c", "i"), math_run(" − "), math_sub("c", "j"), math_run("||")], "2"), math_run("/(2"), math_sup("σ", "2"), math_run(")](1 + λ"), math_sub("IoU", "ij"), math_run(")")], 4)
    add_para(document, "式中，距离项由两个候选的世界坐标质心确定，σ 控制距离衰减速度，IoU 表示世界坐标包围框的交并比，λ 为重叠项权重。分组过程中实施“一切片一检测”约束：在同一个融合组内，每个来源切片最多保留一个检测。融合对象保留来源切片、原始检测和代表性掩膜，以支撑后续点云裁取和逐石复核。")

    add_heading(document, "2.3.3 点云辅助的三维几何筛查", 3)
    add_para(document, "融合对象的局部点云采用“包围框粗筛—掩膜投影精筛”两步提取。系统先以融合包围框在 XY 空间索引中查询候选点，再仅保留落入来源掩膜投影范围的点。该处理避免逐石扫描全场点云，同时使二维岩块轮廓与三维点集在同一空间参考下对应。该步骤对应图 2-2 中的点云几何筛查环节。")
    add_para(document, "由于摄影测量点云主要描述可见上表面，本文从全场点云建立 GroundDEM 作为局部高度参考。对于候选岩块点集中的第 k 个点，其相对于 GroundDEM 的高度定义为：")
    add_formula(document, [math_sub("h", "k"), math_run(" = "), math_sub("z", "k"), math_run(" − "), math_sub("z", "g"), math_run("("), math_sub("x", "k"), math_run(", "), math_sub("y", "k"), math_run(")")], 5)
    add_para(document, "三维筛查同时考虑局部点数、高程极差、相对高度的高位分位数和抬升点比例。该步骤用于识别缺乏明显空间起伏或点云支持不足的二维候选。筛查接受率仅描述当前点云条件下的几何支持情况，不能替代基于人工标注得到的 Precision、Recall 或 F1-score。")

    add_heading(document, "2.3.4 GroundDEM 参考的 2.5D 体积估计", 3)
    add_para(document, "通过三维筛查的对象进入体积计算。摄影测量点云缺少岩块底部与遮挡面的完整描述，因而不宜直接闭合为实体。本文在岩块水平范围内建立局部网格，以可见上表面相对于 GroundDEM 的正高度进行积分。该步骤对应图 2-2 中的 2.5D 体积估计环节。对于网格单元 i，有效高度和 2.5D 体积写为：")
    add_formula(document, [math_sub("h", "i"), math_run(" = max("), math_sup_group([math_sub("z", "i")], "top"), math_run(" − "), math_sup_group([math_sub("z", "i")], "ground"), math_run(", 0),     "), math_sub("V", "2.5D"), math_run(" = Σ "), math_sub("h", "i"), math_sup("Δ", "2")], 6)
    add_para(document, "体积结果经过独立的数值质量控制后进入统计。为构建不使用高程信息的对照指标，本文还由融合等效粒径计算二维等效球代理体积。二者在同一岩块集合上的比较用于描述二维形状假设与局部起伏信息引入后的相对差异；缺少逐石体积真值时，该比较不构成任何一种估计的绝对准确度验证。")

    add_heading(document, "2.4 实施参数与试验设计", 2)
    add_para(document, "矿区 A 用于确定流程参数，矿区 B 采用同一模型权重和同一参数执行。表 2-2 汇总正文需要报告的实施设置。切片尺度、融合阈值和体积网格尺度并非普适常数，应通过矿区 A 的预试验或敏感性分析说明其取值依据；一旦确定，所有参数在矿区 B 的独立运行中保持不变。")
    add_table(document, "表 2-2 流程实施参数", ["处理环节", "设置", "本文采用的数值"], [
        ["DOM 自适应切片", "初始/最小切片边长；重叠宽度；边缘密度阈值；有效内容比例", "20 m / 10 m；0.5 m；0.08；0.05"],
        ["实例分割与二维量测", "输入尺寸；置信度阈值；最小报告粒径", "1024；0.35；0.5 m"],
        ["跨切片融合", "最大关联距离；距离尺度；最低 IoU；正关联阈值", "3.5 m；0.7 m；0.15；0.45"],
        ["三维几何筛查", "最少点数；最小高程极差；P90 相对高度；最小抬升比例", "60；0.18 m；0.12 m；0.20"],
        ["GroundDEM 与体积估计", "GroundDEM 网格；局部积分网格；质量控制最少点数；质量控制高程极差", "0.5 m；0.05 m；30；0.08 m"],
    ], [3.4, 6.4, 6.8])
    add_para(document, "结果评价分为两个层次。首先记录从切片、二维候选、跨切片融合、三维筛查到体积质量控制的对象数量变化，以说明测量流程的可执行性和数据流转关系。其次，在矿区 B 使用冻结配置重复该流程，并比较两个矿区中岩块粒径和体积统计的可获得性。三维筛查的通过与拒绝数量不作为检测准确率；二维代理体积与 2.5D 体积的差异不作为绝对体积误差。若补充人工检测标注或独立体积参考，可在此基础上进一步报告相应精度指标。")

    add_heading(document, "3 结果", 1)
    add_para(document, "本章围绕流程是否在两个矿区中以同一口径完成执行、三维筛查如何改变候选集合，以及二维和 2.5D 测量结果呈现何种差异展开。所有数值、图表和代表性案例均应来自同一冻结实验版本。")
    add_heading(document, "3.1 两个矿区的流程执行与对象数量变化", 2)
    add_heading(document, "3.2 三维几何筛查的对象变化与代表性案例", 2)
    add_heading(document, "3.3 岩块粒径与 GroundDEM 参考的 2.5D 体积统计", 2)
    add_heading(document, "3.4 二维代理体积与 2.5D 估计的相对比较", 2)
    add_heading(document, "3.5 独立矿区运行的结果归纳", 2)
    try:
        document.save(OUTPUT)
        return OUTPUT
    except PermissionError:
        document.save(LOCKED_OUTPUT_FALLBACK)
        return LOCKED_OUTPUT_FALLBACK


if __name__ == "__main__":
    print(build_document())
