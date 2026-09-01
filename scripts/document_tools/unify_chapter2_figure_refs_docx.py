from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PAPER_ROOT = Path(
    r"C:\Users\Administrator\WPSDrive\1714584739\WPS企业云盘\杭州电子科技大学\我的企业文档\Measurement期刊"
)
DOCX_PATH = PAPER_ROOT / "中文论文第二章_材料与方法与第三章结果框架.docx"
FALLBACK_PATH = PAPER_ROOT / "中文论文第二章_材料与方法与第三章结果框架_Fig2图号统一版.docx"
ROOT = Path(__file__).resolve().parents[1]
FIG_2_2_IMAGE = ROOT / "experiments" / "visualization" / "outputs" / "FIG-3-1" / "fig_3_1.png"


def set_body_font(paragraph) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def replace_text(paragraph, text: str, *, centered: bool = False) -> None:
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_body_font(paragraph)


def drawing_count(doc: Document) -> int:
    return sum(
        1
        for paragraph in doc.paragraphs
        for run in paragraph.runs
        if run._element.xpath(".//w:drawing")
    )


def insert_fig_2_2_if_missing(doc: Document) -> bool:
    if drawing_count(doc) >= 2:
        return False
    if not FIG_2_2_IMAGE.exists():
        raise FileNotFoundError(FIG_2_2_IMAGE)

    caption = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "图 2-2 DOM 与点云协同岩块测量流程":
            caption = paragraph
            break
    if caption is None:
        raise RuntimeError("未找到图 2-2 图题，无法定位图片插入位置。")

    figure_paragraph = doc.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.add_run().add_picture(str(FIG_2_2_IMAGE), width=Cm(15.8))
    caption._p.addprevious(figure_paragraph._p)
    return True


def main() -> None:
    doc = Document(DOCX_PATH)
    changed: list[int] = []
    to_delete = []

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()

        if text.startswith("矿区 A 用于流程开发、参数确定和冻结；矿区 B 在不改变模型权重与处理参数的条件下用于独立运行。") and "图 2-1" not in text:
            replace_text(
                paragraph,
                "矿区 A 用于流程开发、参数确定和冻结；矿区 B 在不改变模型权重与处理参数的条件下用于独立运行。"
                "图 2-1 给出了两个研究区的输入数据对照，其中图 2-1(a,b) 为矿区 A 的 DOM 与同一坐标窗口内摄影测量点云，"
                "图 2-1(c,d) 为矿区 B 的 DOM 与同一坐标窗口内摄影测量点云。"
                "两类数据进入计算前均进行空间一致性检查：DOM 的世界文件用于建立像素与投影坐标的关系，点云需转换至同一平面参考框架。"
                "该检查是后续掩膜回投、跨切片融合和局部点云裁取能够对应同一岩块的前提。",
            )
            changed.append(idx)

        elif text == "图 2-1 两个研究区 DOM 与同源摄影测量点云的空间对应关系":
            replace_text(
                paragraph,
                "图 2-1 两个研究区 DOM 与摄影测量点云输入对照",
                centered=True,
            )
            changed.append(idx)

        elif text.startswith("测量流程由 DOM 自适应切片、实例分割与二维量测、跨切片重复检测融合、点云几何筛查及 GroundDEM 参考的 2.5D 体积估计组成。") and "图 2-2" not in text:
            replace_text(
                paragraph,
                "测量流程由 DOM 自适应切片、实例分割与二维量测、跨切片重复检测融合、点云几何筛查及 GroundDEM 参考的 2.5D 体积估计组成。"
                "各阶段以世界坐标为共同索引，最终为每个岩块输出可追溯的二维属性、三维筛查状态和体积估计。"
                "图 2-2 概括了从输入数据到逐石量测结果的主线流程。",
            )
            changed.append(idx)

        elif text.startswith("图 2-3 DOM-点云协同测量中的关键处理步骤示例"):
            to_delete.append((idx, paragraph))

        elif text.startswith("切片重叠保留了边界附近岩块的完整信息"):
            replace_text(
                paragraph,
                "切片重叠保留了边界附近岩块的完整信息，也使同一岩块可能在不同切片中产生重复检测。"
                "为避免重复计数，掩膜质心和包围框先由式（1）转换至世界坐标，再在来自不同切片的检测之间建立关联。"
                "关联同时考虑质心距离、包围框交并比及来源切片约束；满足条件的候选通过相关聚类合并为单一岩块对象。"
                "该过程对应图 2-2 中的跨切片重复检测融合环节。",
            )
            changed.append(idx)

        elif text.startswith("融合对象的局部点云采用“包围框粗筛—掩膜投影精筛”两步提取。"):
            replace_text(
                paragraph,
                "融合对象的局部点云采用“包围框粗筛—掩膜投影精筛”两步提取。"
                "系统先以融合包围框在 XY 空间索引中查询候选点，再仅保留落入来源掩膜投影范围的点。"
                "该处理避免逐石扫描全场点云，同时使二维岩块轮廓与三维点集在同一空间参考下对应。"
                "该步骤对应图 2-2 中的点云几何筛查环节。",
            )
            changed.append(idx)

        elif text.startswith("通过三维筛查的对象进入体积计算。本文采用 GroundDEM 参考的 2.5D 表示"):
            replace_text(
                paragraph,
                "通过三维筛查的对象进入体积计算。本文采用 GroundDEM 参考的 2.5D 表示，"
                "在岩块水平范围内建立局部网格，并以可见表面相对于 GroundDEM 的正高度进行积分。"
                "该步骤对应图 2-2 中的 2.5D 体积估计环节。对于网格单元 i，有效高度和 2.5D 体积写为：",
            )
            changed.append(idx)

    for _, paragraph in reversed(to_delete):
        delete_paragraph(paragraph)

    inserted_fig_2_2 = insert_fig_2_2_if_missing(doc)

    try:
        doc.save(DOCX_PATH)
        output = DOCX_PATH
    except PermissionError:
        doc.save(FALLBACK_PATH)
        output = FALLBACK_PATH

    print(output)
    print("changed", changed)
    print("deleted", [idx for idx, _ in to_delete])
    print("inserted_fig_2_2", inserted_fig_2_2)


if __name__ == "__main__":
    main()
