from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from build_chinese_paper_framework import (
    add_formula,
    add_heading,
    add_para,
    add_placeholder,
    add_rich_para,
    math_fraction,
    math_radical,
    math_run,
    math_sub,
    math_sup,
    math_sup_group,
    remove_existing_body,
)


PAPER_ROOT = Path(
    r"C:\Users\Administrator\WPSDrive\1714584739\WPS企业云盘\杭州电子科技大学\我的企业文档\Measurement期刊"
)
SOURCE = PAPER_ROOT / "中文论文框架_修订版.docx"
OUTPUT = PAPER_ROOT / "中文论文第三章_方法重构稿.docx"


def build_document():
    document = Document(SOURCE)
    remove_existing_body(document)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    add_heading(document, "3 DOM 与点云协同的岩块测量方法", 1)

    add_heading(document, "3.1 测量问题与输入数据要求", 2)
    add_para(
        document,
        "本文研究对象为露天矿岩石堆积场景中的可见岩块，目标是从大范围数字正射影像中获得具有明确空间位置和二维尺度的候选实例，并结合与其同源的摄影测量点云完成三维几何筛查和体积估计。方法输入包括带地理参考的数字正射影像（digital orthophoto map, DOM）以及由同一倾斜摄影重建成果导出的表面点云。前者提供纹理、轮廓和投影尺度，后者提供可见表面的三维坐标。",
    )
    add_para(
        document,
        "DOM 与点云必须位于相容的空间参考体系中。对于每个研究场景，影像像素坐标通过世界文件或等价的仿射参数转换为投影平面坐标；点云则保留重建成果中的世界坐标。由此，切片中的局部掩膜可以恢复到全幅 DOM，并进一步与点云中的水平位置建立对应。该条件是跨切片融合、局部点云裁取和逐石量测能够连续衔接的基础。",
    )
    add_formula(document, [math_run("x = C + Au + Bv,     y = F + Du + Ev")], 1)
    add_para(
        document,
        "式中，(u, v) 为像素坐标，(x, y) 为对应的世界坐标；A 和 E 表示两个方向的像元尺度，B 和 D 为旋转项，C 和 F 为影像原点的投影坐标。不同矿区可使用不同的投影坐标系和空间分辨率，但同一场景内的 DOM、世界文件和点云必须保持一致，不依赖未经说明的经验平移。",
    )
    add_para(
        document,
        "经过完整处理后，每个岩块对象包含来源切片与检测实例、世界坐标位置、二维投影面积、等效粒径、融合成员、局部点云、三维筛查状态和体积估计。摄影测量点云通常只描述可见上表面，不能假设其天然构成封闭实体，因此本文采用地面参考的 2.5D 积分，而不把表面点云直接闭合形成的凸包作为主要测量模型。",
    )
    add_placeholder(
        document,
        "此处插入图 3-1",
        "图 3-1 DOM、世界坐标与同源摄影测量点云的空间对应关系",
    )

    add_heading(document, "3.2 方法总体框架", 2)
    add_para(
        document,
        "所提方法由二维候选生成、跨切片实例融合、点云几何筛查和体积估计四个连续环节构成。首先，根据 DOM 的局部边缘分布组织模型输入，使大范围影像能够在保留原始地面分辨率的条件下分块推理；实例分割模型随后输出岩块掩膜，并据此计算投影面积、等效粒径和世界坐标属性。",
    )
    add_para(
        document,
        "相邻切片之间的重叠会使同一岩块形成多个检测结果。方法在世界坐标中建立跨切片关联图，通过距离、包围框重叠和来源切片约束生成单一融合对象。融合掩膜随后投影至点云平面范围，采用空间索引和掩膜多边形提取局部点集，并基于相对 GroundDEM 的高程特征判断候选是否具有足够的三维几何支持。",
    )
    add_para(
        document,
        "通过三维筛查的对象进入 GroundDEM 参考的 2.5D 网格积分，同时计算仅依赖等效粒径的二维代理体积。第三章定义各参数在算法中的含义与作用；具体取值、参数确定过程以及多矿区统一执行协议在实验设计章节给出。用于独立验证的矿区不参与参数调整，以区分方法开发与跨场景测试。",
    )
    add_placeholder(
        document,
        "此处插入图 3-2",
        "图 3-2 DOM 与点云协同岩块测量方法的总体流程",
    )

    add_heading(document, "3.3 DOM 候选生成与二维量测", 2)
    add_heading(document, "3.3.1 边缘密度引导的四叉树切片", 3)
    add_para(
        document,
        "大范围高分辨率 DOM 的像素尺寸通常超过实例分割网络的直接输入范围。若将整幅影像缩放至模型输入尺寸，岩块边界和小尺度纹理会被压缩；若采用统一的小窗口遍历，则低纹理区域和影像外部背景也会产生不必要的推理开销。为此，本文根据局部边缘密度递归划分影像区域，使切片尺度随局部结构复杂度变化。",
    )
    add_para(
        document,
        "DOM 首先转换为灰度影像，并利用边缘算子获得二值边缘图。对于候选区域 R，边缘密度定义为区域内边缘像素数占总像素数的比例：",
    )
    add_formula(
        document,
        [
            math_sub("ρ", "e"),
            math_run("(R) = "),
            math_fraction(
                [math_sub("N", "e"), math_run("(R)")],
                [math_sub("N", "p"), math_run("(R)")],
            ),
        ],
        2,
    )
    add_rich_para(
        document,
        [
            "其中，",
            ("sub", "N", "e"),
            "(R) 为区域内边缘像素数，",
            ("sub", "N", "p"),
            "(R) 为区域总像素数。考虑到 DOM 外部可能包含近黑色无效区域，进一步定义有效内容比例：",
        ],
    )
    add_formula(
        document,
        [
            math_sub("ρ", "c"),
            math_run("(R) = "),
            math_fraction(
                [math_run("N["), math_sub("I", "g"), math_run("(u,v) > "), math_sub("T", "b"), math_run("]")],
                [math_sub("N", "p"), math_run("(R)")],
            ),
        ],
        3,
    )
    add_rich_para(
        document,
        [
            "式中，",
            ("sub", "I", "g"),
            "(u,v) 为灰度值，",
            ("sub", "T", "b"),
            " 为近黑背景判定阈值。当有效内容比例低于下限或区域中不存在边缘时，该区域不进入推理；否则，若边缘密度达到分裂阈值且当前边长仍大于允许的最小切片尺度，则将其划分为四个子区域。完成递归划分后，在切片边界增加重叠带，以减轻岩块被截断的问题。背景阈值、内容比例阈值、边缘密度阈值、初始与最小切片尺度及重叠宽度均作为实验参数统一设置，而不在方法定义中解释为普适最优值。",
        ],
    )
    add_para(
        document,
        "每个保留切片记录其全局像素范围、世界坐标边界、边缘密度和内容比例。局部检测结果因而能够恢复到全幅 DOM，而重叠区域形成的重复观测则在后续世界坐标融合阶段统一处理。",
    )
    add_placeholder(
        document,
        "此处插入图 3-3",
        "图 3-3 边缘密度引导的四叉树切片原理",
    )

    add_heading(document, "3.3.2 实例分割与二维几何属性计算", 3)
    add_para(
        document,
        "保留切片输入 YOLO11m-seg 实例分割模型，模型对每个候选输出类别置信度、包围框和二值掩膜。本文使用掩膜而非矩形包围框计算岩块尺度，因为掩膜能够保留不规则俯视轮廓。模型输入尺寸、置信度阈值、单切片最大检测数和推理尺度属于实验配置，在多矿区测试中按照预先确定的统一协议执行。",
    )
    add_rich_para(
        document,
        [
            "设实例掩膜的前景像素数为 ",
            ("sub", "n", "p"),
            "，DOM 在横、纵方向的地面采样距离分别为 ",
            ("sub", "s", "x"),
            " 和 ",
            ("sub", "s", "y"),
            "，则二维投影面积为",
        ],
    )
    add_formula(
        document,
        [
            math_sub("A", "2D"),
            math_run(" = "),
            math_sub("n", "p"),
            math_sub("s", "x"),
            math_sub("s", "y"),
        ],
        4,
    )
    add_para(
        document,
        "为以统一的一维尺度描述不同形状的投影轮廓，定义与该投影面积相等的圆直径为等效粒径：",
    )
    add_formula(
        document,
        [
            math_sub("d", "eq"),
            math_run(" = "),
            math_radical(
                [
                    math_fraction(
                        [math_run("4"), math_sub("A", "2D")],
                        [math_run("π")],
                    )
                ]
            ),
        ],
        5,
    )
    add_rich_para(
        document,
        [
            "等效粒径仅表示与实例掩膜具有相同投影面积的圆直径，不等同于岩块的真实长轴、短轴或三维粒径。为明确论文统计的尺度范围，可设置最小报告粒径 ",
            ("sub", "d", "min"),
            "；低于该阈值的候选不进入后续主分析。该阈值描述报告协议，而非模型理论上的最小检测能力。每个保留候选同时记录来源切片、置信度、全局像素位置、世界坐标质心、世界坐标包围框、投影面积和等效粒径。",
        ],
    )
    add_placeholder(
        document,
        "此处插入图 3-4",
        "图 3-4 实例掩膜的世界坐标恢复与二维几何量测",
    )

    add_heading(document, "3.4 跨切片重复检测融合", 2)
    add_para(
        document,
        "切片重叠能够保留边缘附近岩块的完整信息，但也会使同一岩块在不同切片中产生多个位置相近、轮廓略有差异的检测结果。若直接累计这些实例，岩块数量和后续体积统计将受到重复计数影响。本文仅在来自不同切片的检测之间建立重复关联，同一切片中的不同实例默认表示不同候选。",
    )
    add_para(
        document,
        "为避免对全部检测进行全连接比较，首先按照世界坐标质心建立规则网格索引，并仅在最大搜索距离内检索候选对。对于候选 i 和 j，定义距离与包围框交并比共同作用的关联权重：",
    )
    add_formula(
        document,
        [
            math_sub("w", "ij"),
            math_run(" = exp[−"),
            math_sup_group(
                [
                    math_run("||"),
                    math_sub("c", "i"),
                    math_run(" − "),
                    math_sub("c", "j"),
                    math_run("||"),
                ],
                "2",
            ),
            math_run("/(2"),
            math_sup("σ", "2"),
            math_run(")](1 + λ"),
            math_sub("IoU", "ij"),
            math_run(")"),
        ],
        6,
    )
    add_rich_para(
        document,
        [
            "式中，",
            ("sub", "c", "i"),
            " 和 ",
            ("sub", "c", "j"),
            " 为世界坐标质心，σ 控制距离衰减速度，",
            ("sub", "IoU", "ij"),
            " 为世界坐标包围框的交并比，λ 为重叠项权重。只有当两个包围框相交、IoU 达到最低要求且关联权重不低于正关联阈值时，二者之间才建立关联边。网格尺度、最大搜索距离、σ、λ、最低 IoU 和正关联阈值的具体取值在实验设计章节中统一给出。",
        ],
    )

    add_para(
        document,
        "在关联图上采用以枢轴（pivot）为中心的分组过程。算法依次选择尚未分组的检测作为枢轴，并将满足正关联条件的有效邻居纳入同一候选组。由于局部过分割可能使同一来源切片产生多个相近检测，分组时增加来源约束：在一个融合组内，每个来源切片最多保留一个检测；若存在多个候选，则优先选择与枢轴关联权重更高者，权重相同时保留置信度更高者。",
    )
    add_para(
        document,
        "每个融合组被赋予唯一岩块标识。融合后的投影面积和等效粒径取成员中位数，以降低单个边界截断实例的影响；融合包围框取成员包围框并集，置信度取成员均值，同时保留全部来源检测索引和代表性检测。该阶段的作用是形成不重复的候选测量对象，融合数量的变化本身不等同于检测精度提升。",
    )
    add_placeholder(
        document,
        "此处插入图 3-5",
        "图 3-5 跨切片候选关联与重复检测融合",
    )

    add_heading(document, "3.5 点云辅助的三维几何筛查", 2)
    add_heading(document, "3.5.1 局部点云提取与 GroundDEM 构建", 3)
    add_para(
        document,
        "为使融合候选与三维表面对应，首先将其来源掩膜边界恢复到全局像素坐标，并依据式（1）转换为世界坐标多边形。局部点云提取采用包围框粗筛与掩膜多边形精筛相结合的方式：先通过 XY 网格索引查询融合包围框及其外扩范围内的候选点，再保留落入任一来源掩膜投影范围的点。",
    )
    add_para(
        document,
        "这一过程既避免了针对每个岩块重复扫描全场点云，也使边缘切片中的多个掩膜能够共同限定融合岩块的水平范围。空间索引网格尺度和包围框外扩距离属于实现参数，应根据点云密度与目标尺度预先设定，并在各测试场景中遵循相同的设置规则。",
    )

    add_para(
        document,
        "摄影测量点云主要描述堆积体可见表面。为获得相对高程，需要从全场点云建立连续的地面参考面 GroundDEM。点云在水平面上划分为规则网格；对于具有足够采样点的网格单元 g，以其中高程的低分位数作为初始参考高程：",
    )
    add_formula(
        document,
        [
            math_sub("z", "g"),
            math_run(" = "),
            math_sub("Q", "q"),
            math_run("({"),
            math_sub("z", "k"),
            math_run(" | ("),
            math_sub("x", "k"),
            math_run(","),
            math_sub("y", "k"),
            math_run(") ∈ g})"),
        ],
        7,
    )
    add_rich_para(
        document,
        [
            "式中，q 为低位分位数，",
            ("sub", "z", "g"),
            " 为网格单元的参考高程。低分位数能够减弱岩块上表面高点对局部地面估计的影响。采样不足或没有直接观测值的网格由邻近有效单元填补，任意位置的地面高程再由邻近网格插值得到。GroundDEM 网格分辨率、分位数、点云下采样步长和每格最小点数在实验设置中报告。",
        ],
    )

    add_heading(document, "3.5.2 三维几何筛查准则", 3)
    add_para(
        document,
        "对于候选岩块点集中的第 k 个点，定义其相对于 GroundDEM 的高度为",
    )
    add_formula(
        document,
        [
            math_sub("h", "k"),
            math_run(" = "),
            math_sub("z", "k"),
            math_run(" − "),
            math_sub("z", "g"),
            math_run("("),
            math_sub("x", "k"),
            math_run(","),
            math_sub("y", "k"),
            math_run(")"),
        ],
        8,
    )
    add_para(
        document,
        "筛查同时考虑局部点数、高程极差、相对高度的高位分位数和抬升点比例。点数用于约束点云支持度，高程极差描述整体起伏，高位分位数反映主体抬升高度，抬升点比例则限制凸起区域在候选投影范围内的覆盖程度。候选只有同时满足预设条件才进入体积估计；未满足条件的对象保留具体失败原因，允许同一候选对应多个原因。",
    )
    add_para(
        document,
        "三维筛查用于识别缺乏明显空间起伏或点云支持不足的二维候选，并不重新完成岩石语义识别。因此，筛查接受率描述的是当前点云条件下的几何支持情况，不能替代基于人工标注计算的 Precision、Recall 或 F1-score。",
    )
    add_placeholder(
        document,
        "此处插入图 3-6",
        "图 3-6 局部点云提取、GroundDEM 构建与三维几何筛查",
    )

    add_heading(document, "3.6 GroundDEM 参考的 2.5D 体积估计", 2)
    add_heading(document, "3.6.1 2.5D 网格积分与质量控制", 3)
    add_para(
        document,
        "通过三维筛查的岩块点集进入体积计算。由于表面点云通常缺少底部和遮挡面，本文不要求点云形成封闭实体，而是在岩块水平范围内建立局部规则网格。对于包含岩块点的第 i 个单元，取单元内最大高程作为可见上表面高程，并在单元中心查询 GroundDEM 高程。单元有效高度定义为",
    )
    add_formula(
        document,
        [
            math_sub("h", "i"),
            math_run(" = max("),
            math_sup_group([math_sub("z", "i")], "top"),
            math_run(" − "),
            math_sup_group([math_sub("z", "i")], "ground"),
            math_run(", 0)"),
        ],
        9,
    )
    add_para(
        document,
        "设局部积分网格边长为 Δ，则岩块的 2.5D 体积估计为",
    )
    add_formula(
        document,
        [
            math_sub("V", "2.5D"),
            math_run(" = Σ "),
            math_sub("h", "i"),
            math_sup("Δ", "2"),
        ],
        10,
    )
    add_para(
        document,
        "该积分将每个有效网格视为相对于局部地面参考面的竖直柱体，只使用可见上表面和 GroundDEM，不对岩块底部形状施加封闭凸包假设。网格尺度越小，局部起伏表达越细，但对点云密度和计算量的要求也越高，因此其具体取值应在实验设置中明确并保持一致。",
    )
    add_para(
        document,
        "体积计算后设置独立的数值质量控制，用于排除点数不足、高程范围过小、地面查询失败、没有正高度网格或体积非正的结果。其目的在于确保进入统计的 2.5D 估计为可计算的有限正值，而不是重复执行三维候选筛查。质量控制阈值应在实验章节中列出，并在不同矿区使用相同定义。",
    )
    add_para(
        document,
        "最终逐石记录包含世界坐标位置、投影面积、等效粒径、来源检测、局部点数、相对高度指标、三维筛查状态、2.5D 体积、二维代理体积和质量控制标记。该对象级输出为后续多矿区实验中的检测评价、融合复核、筛查分析、粒径分组和体积比较提供统一的数据基础。",
    )

    add_heading(document, "3.6.2 二维等效球代理体积（对照指标）", 3)
    add_para(
        document,
        "为在不使用点云高程信息的条件下构建可复现的对照指标，根据融合后的等效粒径计算二维等效球代理体积：",
    )
    add_formula(
        document,
        [
            math_sub("V", "2D"),
            math_run(" = "),
            math_fraction([math_run("π")], [math_run("6")]),
            math_sup_group([math_sub("d", "eq")], "3"),
        ],
        11,
    )
    add_para(
        document,
        "该指标假设岩块为直径等于二维等效粒径的球体，因而不包含任何实际高度信息。设置该对照项的目的，是在同一岩块集合上量化二维尺度假设与 GroundDEM 参考 2.5D 估计之间的差异，并分析这种差异在不同粒径分组中的表现；它不用于校正、验证或替代 2.5D 体积估计。缺少独立实体体积参考时，两者的相关性、比值和分组差异仅表示相对差异，不能据此判断任何一种体积估计的绝对准确性。",
    )

    add_placeholder(
        document,
        "此处插入图 3-7",
        "图 3-7 GroundDEM 参考的 2.5D 网格积分与二维代理体积",
    )

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
