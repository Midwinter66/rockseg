from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"C:\Users\Administrator\WPSDrive\1714584739\WPS企业云盘\杭州电子科技大学\我的企业文档\Measurement期刊")
SOURCE = ROOT / "文字文稿.docx"
OUTPUT = ROOT / "中文论文框架.docx"


def remove_existing_body(document):
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def add_para(document, text, indent=True):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return p


def add_rich_para(document, parts, indent=True):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    for part in parts:
        if isinstance(part, str):
            run = p.add_run(part)
        else:
            kind, base, *scripts = part
            run = p.add_run(base)
            script_run = p.add_run(scripts[0])
            if kind == "sub":
                script_run.font.subscript = True
            elif kind == "sup":
                script_run.font.superscript = True
            elif kind == "subsup":
                script_run.font.subscript = True
                super_run = p.add_run(scripts[1])
                super_run.font.superscript = True
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
        if not isinstance(part, str):
            script_run.font.name = "Times New Roman"
            script_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            script_run.font.size = Pt(8)
            if kind == "subsup":
                super_run.font.name = "Times New Roman"
                super_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                super_run.font.size = Pt(8)
    return p


def math_run(text):
    run = OxmlElement("m:r")
    run_props = OxmlElement("m:rPr")
    font = OxmlElement("m:font")
    font.set(qn("m:val"), "Cambria Math")
    run_props.append(font)
    run.append(run_props)
    text_element = OxmlElement("m:t")
    text_element.text = text
    run.append(text_element)
    return run


def math_sub(base, subscript):
    element = OxmlElement("m:sSub")
    base_element = OxmlElement("m:e")
    base_element.append(math_run(base))
    sub_element = OxmlElement("m:sub")
    sub_element.append(math_run(subscript))
    element.extend((base_element, sub_element))
    return element


def math_sup(base, superscript):
    element = OxmlElement("m:sSup")
    base_element = OxmlElement("m:e")
    base_element.append(math_run(base))
    sup_element = OxmlElement("m:sup")
    sup_element.append(math_run(superscript))
    element.extend((base_element, sup_element))
    return element


def math_sup_element(base, superscript):
    element = OxmlElement("m:sSup")
    base_element = OxmlElement("m:e")
    base_element.append(base)
    sup_element = OxmlElement("m:sup")
    sup_element.append(math_run(superscript))
    element.extend((base_element, sup_element))
    return element


def math_sup_group(items, superscript):
    element = OxmlElement("m:sSup")
    base_element = OxmlElement("m:e")
    for item in items:
        base_element.append(item)
    sup_element = OxmlElement("m:sup")
    sup_element.append(math_run(superscript))
    element.extend((base_element, sup_element))
    return element


def math_fraction(numerator, denominator):
    element = OxmlElement("m:f")
    num_element = OxmlElement("m:num")
    den_element = OxmlElement("m:den")
    for item in numerator:
        num_element.append(item)
    for item in denominator:
        den_element.append(item)
    element.extend((num_element, den_element))
    return element


def math_radical(items):
    element = OxmlElement("m:rad")
    rad_props = OxmlElement("m:radPr")
    degree_hide = OxmlElement("m:degHide")
    degree_hide.set(qn("m:val"), "1")
    rad_props.append(degree_hide)
    degree = OxmlElement("m:deg")
    degree.append(math_run(""))
    base = OxmlElement("m:e")
    for item in items:
        base.append(item)
    element.extend((rad_props, degree, base))
    return element


def add_formula(document, items, number):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    equation = OxmlElement("m:oMath")
    for item in items:
        equation.append(item)
    p._p.append(equation)
    run = p.add_run(f"    ({number})")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    return p


def add_placeholder(document, label, caption):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[{label}]")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(caption)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)


def add_table(document, title, headers, rows, widths=None):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10)

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(document, text, level):
    p = document.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    return p


def build_document():
    document = Document(SOURCE)
    remove_existing_body(document)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    add_heading(document, "3 材料与方法", 1)

    add_heading(document, "3.1 研究场景、输入数据与空间参考", 2)
    add_para(document, "本研究面向露天矿岩石堆积场景的块度与体积统计，输入为同一倾斜摄影 OSGB 重建成果导出的数字正射影像（digital orthophoto map, DOM）和表面点云。DOM 提供连续的纹理、边界与平面尺度信息，点云提供同一场景可见表面的三维坐标。两类数据并非来自独立传感器，而是在统一摄影测量重建与投影坐标框架下对同一对象的二维和三维表达。")
    add_para(document, "DOM 以 GeoTIFF 格式保存，影像尺寸为 8783 × 21713 pixels，空间分辨率为 0.01 m/pixel，平面坐标系统为 CGCS2000 / 3-degree Gauss-Kruger CM 81E（EPSG:4536），覆盖面积约为 19070.53 m²。点云以两个空间相邻的 LAZ 数据块保存，总点数为 146,721,392。本文将这两个数据块作为同一研究场景的组成部分共同参与局部点云提取、三维支持筛查和体积估计，而不将其视为相互独立的研究区域。")
    add_para(document, "像素坐标与世界坐标之间的关系由 DOM 配套的世界文件确定。对于像素坐标 (u, v)，其世界坐标 (x, y) 由六参数仿射变换计算，其中 A 和 E 为像元尺度，B 和 D 为旋转项，C 和 F 为影像原点的投影坐标。当前数据的旋转项为零，横、纵向像元尺度分别为 0.01 m 和 -0.01 m。切片定位、实例掩膜回投、跨切片融合和点云裁剪均使用该统一坐标关系；点云处理采用绝对世界坐标，不额外施加经验 XY 平移。")
    add_formula(document, [math_run("x = C + Au + Bv,     y = F + Du + Ev")], 1)
    add_para(document, "当前场景不包含逐岩块人工体积真值。因此，本文采用由 DOM 候选识别、跨切片重复消解、点云三维筛查和 GroundDEM 约束体积估计构成的测量方法。三维筛查用于判断二维候选是否具有可辨识的几何支持；2.5D 计算用于给出相对于局部地面参考面的体积估计。两者均不构成绝对体积精度验证。")
    add_table(document, "表 3-1 当前研究场景的输入数据与空间参考", ["Item", "Description"], [
        ["DOM", "OSGB-derived GeoTIFF; 8783 × 21713 pixels; 0.01 m/pixel"],
        ["Point cloud", "LAZ surface point cloud derived from the same OSGB model; two adjacent blocks; 146,721,392 points"],
        ["Coordinate reference", "CGCS2000 / 3-degree Gauss-Kruger CM 81E (EPSG:4536)"],
        ["Scene area", "19070.53 m²"],
        ["Coordinate mapping", "Affine DOM-to-world mapping; absolute world coordinates; zero XY shift"],
    ], [4.2, 12.8])
    add_placeholder(document, "此处插入图 3-1", "图 3-1 研究场景及 OSGB 导出 DOM 和点云的对应关系")

    add_heading(document, "3.2 测量方法总体框架", 2)
    add_para(document, "本文的测量链由二维候选生成和三维量测约束两部分组成。首先，依据局部纹理复杂度将大幅 DOM 组织为可推理的影像 tile；随后，YOLO11m-seg 在有效 tile 上输出岩石实例掩膜、检测框与置信度。每个掩膜被转换为实际单位下的投影面积、等效粒径和世界坐标属性，并作为后续融合和点云处理的基础对象。")
    add_para(document, "由于相邻 tile 存在重叠，同一岩块可能产生多次检测。本文在世界坐标中建立跨 tile 关联，并采用相关聚类将潜在重复检测合并为融合候选。随后，依据来源实例掩膜在点云中裁取局部表面点集，并以局部点数、高程范围、相对 GroundDEM 高度和抬升点比例进行三维筛查。")
    add_para(document, "通过三维筛查的对象进入 GroundDEM 参考的 2.5D 网格积分，并同步计算由二维等效粒径构造的等效球代理体积。两种估计量对应同一融合对象，可用于比较二维尺度假设与引入局部高程信息后的估计差异。")
    add_table(document, "表 3-2 所提方法的关键参数设置", ["Stage", "Parameter", "Value"], [
        ["Tiling", "Base / minimum tile size", "10 m / 5 m"],
        ["Tiling", "Canny low / high thresholds", "30 / 90"],
        ["Tiling", "Minimum edge density / overlap / minimum content ratio", "0.10 / 0.5 m / 0.05"],
        ["Detection", "Model / input / confidence", "YOLO11m-seg / 1024 pixels / 0.35"],
        ["Detection", "Maximum detections per tile", "1000"],
        ["Detection", "Minimum equivalent diameter", "0.5 m"],
        ["Fusion", "Spatial-index cell / maximum centroid distance", "1.0 m / 3.5 m"],
        ["Fusion", "Distance scale / positive-association threshold", "0.7 m / 0.45"],
        ["Fusion", "Minimum bbox IoU", "0.15"],
        ["Fusion", "One retained detection per source tile within a fusion group", "Enabled"],
        ["3D screening", "Minimum points / z range", "60 / 0.18 m"],
        ["3D screening", "P90 height / elevated height / elevated ratio", "0.12 m / 0.08 m / 0.20"],
        ["Point-cloud crop", "Bounding-box padding / XY-index cell", "0.5 m / 1.0 m"],
        ["GroundDEM", "Resolution / elevation percentile", "0.5 m / 5th percentile"],
        ["GroundDEM", "Subsample step / minimum points per cell", "100 / 3"],
        ["Volume", "Local integration grid", "0.05 m"],
    ], [3.2, 6.4, 7.4])
    add_placeholder(document, "此处插入图 3-2", "图 3-2 DOM 与点云协同岩块测量方法的总体流程")

    add_heading(document, "3.3 DOM 切片、实例分割与二维量测", 2)
    add_heading(document, "3.3.1 边缘密度引导的四叉树切片", 3)
    add_para(document, "原始 DOM 的空间范围明显大于实例分割模型的直接输入尺度。整幅缩放会压缩岩块边界与小尺度纹理，而完全按原始尺寸推理的计算代价较高。本文采用边缘密度引导的四叉树切片，将计算资源优先分配给纹理和边界变化较多的区域。固定滑窗切片可作为分块推理的参考方案，但本文不将 tile 数量差异直接解释为检测性能差异。")
    add_para(document, "具体而言，DOM 首先转换为灰度影像，并使用低、高阈值分别为 30 和 90 的 Canny 算子提取边缘。对于候选区域 R，边缘密度定义为区域内边缘像素数与总像素数之比，其中 Nₑ(R) 为边缘像素数，Nₚ(R) 为区域总像素数。为排除影像外部的无效黑色背景，进一步以灰度值大于 5 的像素所占比例表示有效内容比例。当该比例低于 0.05 或区域不含边缘像素时，区域被标记为跳过；当边缘密度不低于 0.10 且区域边长大于 5 m 时，区域继续四分。")
    add_formula(document, [math_sub("ρ", "e"), math_run("(R) = "), math_fraction([math_sub("N", "e"), math_run("(R)")], [math_sub("N", "p"), math_run("(R)")])], 2)
    add_para(document, "初始 tile 边长为 10 m，最小 tile 边长为 5 m，相邻切片设置 0.5 m 重叠宽度。四叉树完成空间划分后再向边界扩展重叠区域，以降低岩块被 tile 边缘截断的概率。每个 tile 均保存像素范围、世界坐标范围、边缘密度、有效内容比例和保留状态，以恢复后续实例的空间位置。")
    add_placeholder(document, "此处插入图 3-3", "图 3-3 DOM 的边缘密度引导四叉树切片")

    add_heading(document, "3.3.2 切片级实例分割与二维几何属性计算", 3)
    add_para(document, "每个保留 tile 采用 YOLO11m-seg 进行单类岩石实例分割，输入尺寸为 1024 pixels，置信度阈值为 0.35，采用单尺度推理，每个 tile 最多保留 1000 个检测。模型输出包括实例掩膜、包围框和置信度；掩膜以游程编码保存，并与来源 tile 标识和全局像素原点关联。tile-level 检测仅作为候选集合，需经后续融合和三维筛查后进入统计。")
    add_rich_para(document, ["对于实例掩膜，设前景像素数为 ", ("sub", "n", "p"), "，DOM 在两个方向的地面采样距离为 ", ("sub", "s", "x"), " 和 ", ("sub", "s", "y"), "，则二维投影面积由掩膜像素数计算。为在不同平面轮廓之间建立统一的一维尺度，采用与投影面积相同的等面积圆直径作为等效粒径。该指标服务于报告尺度筛选、粒径分组和二维代理体积计算，不代表岩块的真实长轴、短轴或三维直径。"])
    add_formula(document, [math_sub("A", "2D"), math_run(" = "), math_sub("n", "p"), math_sub("s", "x"), math_sub("s", "y")], 3)
    add_formula(document, [math_sub("d", "eq"), math_run(" = "), math_radical([math_fraction([math_run("4"), math_sub("A", "2D")], [math_run("π")])])], 4)
    add_rich_para(document, ["本文主分析仅保留 ", ("sub", "d", "eq"), " 不小于 0.5 m 的候选。该阈值定义的是本研究当前报告的粒径范围，而非模型理论上的最小可识别尺度。改变该阈值会同时改变检测候选数、融合关系、三维筛查对象和体积统计集合，因此后续敏感性分析必须沿用完整处理链并重新生成所有统计结果。"])
    add_placeholder(document, "此处插入图 3-4", "图 3-4 由实例掩膜到二维几何量测的转换")

    add_heading(document, "3.4 跨切片检测结果的世界坐标融合", 2)
    add_para(document, "重叠切片使同一岩块在不同 tile 中可能出现位置相近但轮廓并不完全一致的检测实例。若将所有 tile-level 实例直接累加，会使数量、面积和体积统计受到重复计数的影响。因此，本文首先将检测结果依据仿射关系回投到世界坐标系，并以候选质心、世界坐标包围框、面积、等效粒径、置信度和来源 tile 构建跨 tile 融合对象。来自同一 tile 的不同实例不参与相互合并。")
    add_para(document, "为减少候选两两比较的计算量，系统以 1.0 m 网格建立质心空间索引，仅在最大质心距离 3.5 m 的邻域内检索跨 tile 候选。对任意两个候选 i 和 j，只有在世界坐标包围框相交且 IoU 不低于 0.15 时，才计算距离和 IoU 联合关联权重。距离项使空间上相近的候选获得更高权重，IoU 项补充包围范围的一致性。")
    add_formula(document, [math_sub("w", "ij"), math_run(" = exp[−"), math_sup_group([math_run("||"), math_sub("c", "i"), math_run(" − "), math_sub("c", "j"), math_run("||")], "2"), math_run("/(2"), math_sup("σ", "2"), math_run(")](1 + λ"), math_sub("IoU", "ij"), math_run(")")], 5)
    add_rich_para(document, ["取 σ = 0.7 m、λ = 0.3，并将 ", ("sub", "w", "ij"), " 不低于 0.45 的候选对定义为正关联。随后采用以 pivot 为中心的相关聚类成组：依次选取未分组候选作为 pivot，并从其正关联邻居中选择可并入同一组的成员。为避免同一来源 tile 内的重复响应被共同纳入一个岩块，聚类时施加“一 tile 一检测”约束；当同一 tile 的多个邻居与 pivot 关联时，仅保留关联权重最高的检测，权重相同时保留置信度更高者。"])
    add_para(document, "每个融合组生成唯一的石块标识。融合面积和等效粒径取成员中位数，融合包围框取成员包围框并集，融合置信度取成员置信度均值；同时保存来源检测索引、来源 tile 编号和代表性检测，以支持后续点云裁取和逐石核查。")
    add_placeholder(document, "此处插入图 3-5", "图 3-5 重叠切片中重复检测结果的世界坐标融合")

    add_heading(document, "3.5 点云辅助的三维几何筛查", 2)
    add_para(document, "融合候选的局部点云采用“包围框粗筛—掩膜投影精筛”的两阶段方式获得。系统首先在 1.0 m XY 网格索引中查询融合包围框及其 0.5 m 外扩范围内的点云候选；再将融合组中来源掩膜的边界转换为世界坐标多边形，仅保留落入任一掩膜投影范围的点。该处理使二维候选与三维点集在同一空间参考下对应，同时避免为每个候选遍历全场点云。")
    add_para(document, "为表征候选点集相对于局部地面的起伏，首先在全场构建 GroundDEM。GroundDEM 的网格分辨率为 0.5 m；每个具有至少 3 个采样点的网格单元取高程第 5 百分位作为初始地面高程，点云在构建前每隔 100 个点进行系统下采样。对缺失网格，以邻近有效单元逐层填补，以保证后续地面高程查询连续。该参考面用于刻画可见表面相对局部低位基准的抬升，而不被视为独立地面真值。")
    add_formula(document, [math_sub("h", "k"), math_run(" = "), math_sub("z", "k"), math_run(" − "), math_sub("z", "g"), math_run("("), math_sub("x", "k"), math_run(", "), math_sub("y", "k"), math_run(")")], 6)
    add_rich_para(document, ["三维筛查同时约束局部点云支持度和相对地面起伏。候选需满足：点数不少于 60；局部高程极差不小于 0.18 m；相对高度 ", ("sub", "h", "k"), " 的第 90 百分位不小于 0.12 m；相对高度不低于 0.08 m 的点所占比例不小于 0.20。未满足任一条件的候选被保留为拒绝记录，并附带具体失败原因。多个失败条件可同时出现，因此按失败原因的计数不应与拒绝候选总数简单相加。"])
    add_para(document, "该步骤用于筛查平坦纹理、阴影边缘或点云支持不足的异常候选。三维筛查反映的是候选在当前点云条件下的几何支持情况，不替代基于人工标注的检测性能评价。")
    add_placeholder(document, "此处插入图 3-6", "图 3-6 融合候选的点云提取与三维几何筛查")

    add_heading(document, "3.6 基于 GroundDEM 的 2.5D 体积估计与二维代理对照", 2)
    add_para(document, "通过三维筛查的候选进入逐石体积计算。OSGB 导出的点云主要描述可见表面，岩块底部通常不形成封闭实体；在此数据条件下，直接以表面壳层构造凸包容易引入不受控的闭合假设。本文因此采用以 GroundDEM 为参考面的 2.5D 网格积分，将局部可见上表面相对于地面参考面的正高度累积为体积估计。")
    add_rich_para(document, ["对于每个岩块，在其局部范围内建立 0.05 m × 0.05 m 水平网格。对含有岩块点的网格单元，取其中最大高程作为上表面高程 ", ("subsup", "z", "i", "top"), "，并在网格中心查询 GroundDEM 高程 ", ("subsup", "z", "i", "ground"), "。单元有效高度取二者之差与零的较大值，逐单元柱体体积累积得到 2.5D 体积。"])
    add_formula(document, [math_sub("h", "i"), math_run(" = max("), math_sup_group([math_sub("z", "i")], "top"), math_run(" − "), math_sup_group([math_sub("z", "i")], "ground"), math_run(", 0),     "), math_sub("V", "2.5D"), math_run(" = Σ "), math_sub("h", "i"), math_sup("Δ", "2")], 7)
    add_para(document, "体积阶段设置质量控制，以排除无法形成数值有效结果的样本。计算结果需满足逐石点数不少于 30、高程极差不小于 0.08 m、体积状态有效且 2.5D 体积为正。该质量控制用于保证积分结果的数值有效性，与前述三维筛查的目的不同。")
    add_rich_para(document, ["为构建不使用高度信息的二维参照，本文以融合后的等效粒径定义等效球代理体积。该代理完全由二维投影尺度导出，因而具有明确的形状假设。将其与 GroundDEM 参考的 2.5D 估计在同一石块集合上比较，可用于分析二维等效球假设与局部起伏信息引入后的相对差异；在缺少逐石体积真值时，这一比较不构成绝对体积准确度验证。"])
    add_formula(document, [math_sub("V", "2D"), math_run(" = "), math_fraction([math_run("π")], [math_run("6")]), math_sup_group([math_sub("d", "eq")], "3")], 8)
    add_placeholder(document, "此处插入图 3-7", "图 3-7 GroundDEM 参考的 2.5D 体积估计与二维代理对照")

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
