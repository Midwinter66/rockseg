from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX_PATH = Path(
    r"C:\Users\Administrator\WPSDrive\1714584739\WPS企业云盘\杭州电子科技大学\我的企业文档\Measurement期刊\中文论文第二章_材料与方法与第三章结果框架.docx"
)


REPLACEMENTS = {
    "为保持实验边界清晰，模型权重、输入尺寸和置信度阈值在矿区 A 完成参数确定后固定，矿区 B 不再据其运行结果调整。投稿前需补充模型训练数据的标注规范、训练集与验证集划分、数据增强、训练轮数、优化器和学习率等复现信息；这些信息用于说明模型来源，不应与本文两个研究区的独立运行数据混写。": (
        "为保持实验边界清晰，模型权重、输入尺寸和置信度阈值在矿区 A 完成参数确定后固定，矿区 B 不再据其运行结果调整。"
        "模型训练数据的标注规范、训练集与验证集划分、数据增强和训练超参数属于模型来源信息；本文的结果分析则限定在两个研究区的冻结流程运行结果，二者在实验叙述中保持区分。"
    ),
    "矿区 A 用于确定流程参数，矿区 B 采用同一模型权重和同一参数执行。表 2 汇总正文需要报告的实施设置。切片尺度、融合阈值和体积网格尺度并非普适常数，应通过矿区 A 的预试验或敏感性分析说明其取值依据；一旦确定，所有参数在矿区 B 的独立运行中保持不变。": (
        "矿区 A 用于确定流程参数，矿区 B 采用同一模型权重和同一参数执行。表 2-2 汇总本文采用的主要实施设置。"
        "切片尺度、融合阈值和体积网格尺度不被视为普适常数，而作为本研究流程的冻结配置记录；一旦确定，所有参数在矿区 B 的独立运行中保持不变。"
    ),
    "结果评价分为两个层次。首先记录从切片、二维候选、跨切片融合、三维筛查到体积质量控制的对象数量变化，以说明测量流程的可执行性和数据流转关系。其次，在矿区 B 使用冻结配置重复该流程，并比较两个矿区中岩块粒径和体积统计的可获得性。三维筛查的通过与拒绝数量不作为检测准确率；二维代理体积与 2.5D 体积的差异不作为绝对体积误差。若补充人工检测标注或独立体积参考，可在此基础上进一步报告相应精度指标。": (
        "基于上述设置，第三章从三个层面报告结果：首先记录切片、二维候选、跨切片融合、三维筛查到体积质量控制的对象数量变化，以说明测量流程的可执行性和数据流转关系；"
        "其次，在矿区 B 使用冻结配置重复该流程，并与矿区 A 的基线输出进行对照；最后比较 2.5D 体积与二维代理体积的差异，以分析引入点云高度信息后的统计变化。"
        "三维筛查的通过与拒绝数量不作为检测准确率；二维代理体积与 2.5D 体积的差异不作为绝对体积误差。"
    ),
    "本章基于前述固定流程，对矿区 A 的基线结果与矿区 B 的独立运行结果进行整理和分析。矿区 B 在不重新训练分割模型、不改变主要处理流程的条件下完成运行，因此其结果主要用于考察该 DOM-点云协同测量流程在另一处矿区场景中的可执行性和输出稳定性。需要说明的是，本章中的三维几何筛查保留率表示候选目标经过点云几何约束后的保留比例，并不等同于人工标注意义上的检测精度；二维代理体积与 2.5D 体积之间的差异用于说明估算方法的差别，也不作为真实体积误差。": (
        "本章按照第二章定义的固定流程，对矿区 A 的基线结果与矿区 B 的独立运行结果进行整理和分析。"
        "矿区 B 在不重新训练分割模型、不改变主要处理参数的条件下完成运行，因此其结果主要用于考察该 DOM-点云协同测量流程在另一处矿区场景中的可执行性和输出稳定性。"
        "本章中的三维几何筛查保留率表示候选目标经过点云几何约束后的保留比例，并不等同于人工标注意义上的检测精度；二维代理体积与 2.5D 体积之间的差异用于说明估算方法差别，也不作为真实体积误差。"
    ),
    "两个矿区均完成了从自适应切片、实例分割、重叠切片融合、三维几何筛查到 2.5D 体积估算的完整流程。矿区 A 的 DOM 覆盖面积为 19070.5279 m²，四叉树切片后保留 98 个有效切片；矿区 B 的 DOM 覆盖面积为 16521.71 m²，保留 125 个有效切片。两处场景的有效切片数量并不完全随 DOM 面积线性变化，这与石料堆积密度、影像边缘分布以及无效背景比例有关。": (
        "两个矿区均完成了从自适应切片、实例分割、重叠切片融合、三维几何筛查到 2.5D 体积估算的完整流程。"
        "矿区 A 的 DOM 覆盖面积为 19070.53 m²，四叉树切片后保留 98 个有效切片；矿区 B 的 DOM 覆盖面积为 16521.71 m²，保留 125 个有效切片。"
        "两处场景的有效切片数量并不完全随 DOM 面积线性变化，这与石料堆积密度、影像边缘分布以及无效背景比例有关。"
    ),
    "通过三维筛查后，矿区 A 有 6929 个对象进入 2.5D 体积统计，矿区 B 有 8538 个对象进入统计。矿区 A 的 2.5D 总体积为 1451.0138 m³，均值为 0.2094 m³，中位数为 0.1151 m³；矿区 B 的 2.5D 总体积为 2090.228 m³，均值为 0.2448 m³，中位数为 0.1279 m³。两处矿区的单石体积均呈现均值高于中位数的特征，说明体积分布受到少量较大石块的影响。": (
        "通过三维筛查后，矿区 A 有 6929 个对象进入 2.5D 体积统计，矿区 B 有 8538 个对象进入统计。"
        "矿区 A 的 2.5D 总体积为 1451.01 m³，均值为 0.209 m³，中位数为 0.115 m³；矿区 B 的 2.5D 总体积为 2090.23 m³，均值为 0.245 m³，中位数为 0.128 m³。"
        "两处矿区的单石体积均呈现均值高于中位数的特征，说明体积分布受到少量较大石块的影响。"
    ),
    "矿区 A 的二维代理总体积为 2221.624 m³，高于对应的 2.5D 总体积 1451.0138 m³；矿区 B 的二维代理总体积为 2916.3142 m³，同样高于 2.5D 总体积 2090.228 m³。在逐石尺度上，矿区 A 中两种体积指标的 Pearson 相关系数为 0.8204，矿区 B 为 0.8182。这说明二维尺度信息与 2.5D 体积之间存在较强相关性，但二维代理方法倾向于将等效直径直接映射为规则几何体，无法反映石块可见表面的真实起伏和局部高度变化。": (
        "矿区 A 的二维代理总体积为 2221.62 m³，高于对应的 2.5D 总体积 1451.01 m³；矿区 B 的二维代理总体积为 2916.31 m³，同样高于 2.5D 总体积 2090.23 m³。"
        "在逐石尺度上，矿区 A 中两种体积指标的 Pearson 相关系数为 0.820，矿区 B 为 0.818。"
        "这说明二维尺度信息与 2.5D 体积之间存在较强相关性，但二维代理方法倾向于将等效直径直接映射为规则几何体，无法反映石块可见表面的真实起伏和局部高度变化。"
    ),
}

CAPTION_REPLACEMENTS = {
    "表 1 研究区输入数据概况": "表 2-1 研究区输入数据概况",
    "图 1 两个研究区 DOM 与同源摄影测量点云的空间对应关系": "图 2-1 两个研究区 DOM 与摄影测量点云输入对照",
    "图 2-1 两个研究区 DOM 与同源摄影测量点云的空间对应关系": "图 2-1 两个研究区 DOM 与摄影测量点云输入对照",
    "图 2 DOM 与点云协同岩块测量流程": "图 2-2 DOM 与点云协同岩块测量流程",
    "表 2 流程实施参数": "表 2-2 流程实施参数",
}


TABLE_CELL_UPDATES = {
    (0, 2, 2): "7269 × 22729 像素\n空间分辨率：0.01 m\n覆盖面积：16521.71 m²\n坐标参考：EPSG:4536",
    (0, 2, 3): "两个 LAZ 数据块\n总点数：187360460\n局部坐标，经范围核验后转换至 DOM 坐标框架",
    (0, 2, 4): "使用冻结流程进行独立运行",
    (2, 1, 2): "19070.53",
    (4, 2, 2): "1451.01",
    (4, 2, 3): "2090.23",
    (4, 3, 2): "0.209",
    (4, 3, 3): "0.245",
    (4, 4, 2): "0.115",
    (4, 4, 3): "0.128",
    (4, 5, 2): "2221.62",
    (4, 5, 3): "2916.31",
    (4, 6, 2): "0.820",
    (4, 6, 3): "0.818",
    (5, 1, 4): "616.82",
    (5, 2, 4): "492.85",
    (5, 3, 4): "574.31",
    (5, 4, 4): "406.24",
    (5, 1, 6): "0.110",
    (5, 2, 6): "0.275",
    (5, 3, 6): "0.630",
    (5, 4, 6): "1.822",
}


def set_font(paragraph, size: float = 10.5, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


def format_paragraph(paragraph) -> None:
    text = paragraph.text.strip()
    style = paragraph.style.name
    if style == "Heading 1":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        set_font(paragraph, 14, True)
    elif style == "Heading 2":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(4)
        set_font(paragraph, 12, True)
    elif style == "Heading 3":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(3)
        set_font(paragraph, 10.5, True)
    else:
        if text.startswith(("表 ", "图 ")):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(3)
        elif text.startswith("[此处插入图"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.5
        set_font(paragraph, 10.5, None)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            edge_data = {"val": "nil"}
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in list(element.attrib):
            del element.attrib[key]
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    rows = table.rows
    for r_idx, row in enumerate(rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.paragraph_format.space_after = Pt(0)
                set_font(paragraph, 9.5, r_idx == 0)

    line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    thick = {"val": "single", "sz": "12", "space": "0", "color": "000000"}
    for cell in rows[0].cells:
        set_cell_border(cell, top=thick, bottom=line)
    for cell in rows[-1].cells:
        # Preserve the top border on the header when the table has only one row.
        if cell in rows[0].cells:
            set_cell_border(cell, top=thick, bottom=thick)
        else:
            set_cell_border(cell, bottom=thick)


def main() -> None:
    doc = Document(DOCX_PATH)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in REPLACEMENTS:
            paragraph.text = REPLACEMENTS[text]
        elif text in CAPTION_REPLACEMENTS:
            paragraph.text = CAPTION_REPLACEMENTS[text]

    for table_idx, row_idx, col_idx in list(TABLE_CELL_UPDATES):
        doc.tables[table_idx].rows[row_idx].cells[col_idx - 1].text = TABLE_CELL_UPDATES[(table_idx, row_idx, col_idx)]

    for paragraph in doc.paragraphs:
        format_paragraph(paragraph)

    for table in doc.tables:
        format_table(table)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
